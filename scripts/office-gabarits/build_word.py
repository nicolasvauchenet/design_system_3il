from pathlib import Path
from copy import deepcopy
from hashlib import sha256
import json, zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT=Path(__file__).resolve().parents[2]
OUT=ROOT/'office/Gabarits'; OUT.mkdir(parents=True,exist_ok=True)
TMP=Path(__file__).parent
REF=ROOT/'office/UML/3iL-UML-Cours.docx'
before=sha256(REF.read_bytes()).hexdigest()
d=Document(REF)
BLUE='005067'; SOFT='EDF4F5'
def el(tag,**attrs):
    x=OxmlElement('w:'+tag)
    for k,v in attrs.items(): x.set(qn('w:'+k),str(v))
    return x
def field(p,code,cache):
    p.add_run()._r.append(el('fldChar',fldCharType='begin'))
    t=el('instrText'); t.set(qn('xml:space'),'preserve'); t.text=' '+code+' '; p.add_run()._r.append(t)
    p.add_run()._r.append(el('fldChar',fldCharType='separate'))
    p.add_run(cache)
    p.add_run()._r.append(el('fldChar',fldCharType='end'))
def para(t='',style=None):return d.add_paragraph(t,style)
def title(t):
    p=para(t,'Heading 1'); p.paragraph_format.page_break_before=True;return p
def heading(t):return para(t,'Heading 2')
def callout(label,text):return para(label+' — '+text,'3iL Repère')

# Retain the real cover, styles, numbering and page system; replace all course content.
cover=list(d.paragraphs[:8]); keep={p._p for p in cover}
for child in list(d._element.body):
    if child not in keep and child.tag!=qn('w:sectPr'): d._element.body.remove(child)
for i,text in {1:'SUPPORT DE COURS',2:'[Titre du cours]',3:'[Sous-titre ou compétence visée]',4:'[Formation] · [Niveau] · [Année universitaire]',5:'[Nom de l’enseignant]\n[Durée du module] · [Version / date]'}.items():
    cover[i].text=text

# Usable style gallery, no outline level for instruction headings.
s=d.styles.add_style('3iL Guide',WD_STYLE_TYPE.PARAGRAPH)
s.base_style=d.styles['Normal'];s.font.name='Arial';s.font.size=Pt(24);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(BLUE)
s.paragraph_format.space_after=Pt(12);s.paragraph_format.keep_with_next=True
for name in ['Normal','Heading 1','Heading 2','Heading 3','3iL Repère','3iL Code','Caption','3iL Guide']:
    d.styles[name].quick_style=True;d.styles[name].priority=1 if name=='Normal' else 2
for name in ['3iL Définition','3iL Exemple','3iL Attention']:
    s=d.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH);s.base_style=d.styles['3iL Repère'];s.quick_style=True;s.priority=3
    if name=='3iL Attention':
        pp=s.element.get_or_add_pPr();pp.append(el('shd',fill='FFF5E9'))
for name in ['3iL Puces','3iL Étapes']:
    s=d.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH);s.base_style=d.styles['Normal'];s.quick_style=True;s.priority=3
    s.paragraph_format.space_after=Pt(5)
    nr=d.part.numbering_part.element
    aid=max(int(x.get(qn('w:abstractNumId'))) for x in nr.findall(qn('w:abstractNum')))+1
    nid=max(int(x.get(qn('w:numId'))) for x in nr.findall(qn('w:num')))+1
    a=el('abstractNum',abstractNumId=aid);a.append(el('multiLevelType',val='singleLevel'))
    lv=el('lvl',ilvl=0);lv.append(el('start',val=1));lv.append(el('numFmt',val='bullet' if name=='3iL Puces' else 'decimal'))
    lv.append(el('lvlText',val='•' if name=='3iL Puces' else '%1.'));lv.append(el('lvlJc',val='left'))
    pp=el('pPr');pp.append(el('ind',left=360,hanging=240));lv.append(pp);a.append(lv);nr.append(a)
    n=el('num',numId=nid);n.append(el('abstractNumId',val=aid));nr.append(n)
    np=el('numPr');np.append(el('ilvl',val=0));np.append(el('numId',val=nid));s.element.get_or_add_pPr().append(np)

p=para('Prise en main · enseignant','3iL Guide')
para('Cette page explique le gabarit. Supprimez-la avant de transmettre le cours aux étudiants.','3iL Repère')
for a,b in [
('1. Créez votre copie','Ouvrez le fichier .dotx par double-clic, puis enregistrez votre cours au format .docx sous un nouveau nom.'),
('2. Remplacez les textes entre crochets','Renseignez la couverture, puis conservez uniquement les pages et blocs utiles à votre cours.'),
('3. Collez sans importer la mise en forme','Lors du collage de votre cours, choisissez « Conserver uniquement le texte ». Appliquez ensuite les styles depuis Accueil > Styles.'),
('4. Utilisez les styles','Titre 1 : chapitre ; Titre 2 : section ; Titre 3 : sous-section ; Normal : texte. Les styles 3iL gèrent les encadrés, listes, étapes et blocs de code.'),
('5. Actualisez et vérifiez','Après les modifications, faites Ctrl+A puis F9. Dans le sommaire, choisissez la mise à jour de toute la table. Vérifiez les pages avant export PDF.')]:
    p=para();p.add_run(a).bold=True;para(b)
callout('Pour gagner du temps','Dupliquez un bloc déjà mis en forme plutôt que de recréer ses réglages. Ne modifiez pas les logos.')
para('Avant diffusion : retirez cette notice, remplacez tous les crochets, relisez les contenus et les sources, vérifiez les tableaux et les légendes. Le corrigé enseignant doit rester dans un fichier séparé.')

p=para('Sommaire','3iL Guide');p.paragraph_format.page_break_before=True
field(para(),'TOC \\o "1-1" \\h \\z','1. [Titre du chapitre] ........................................ 4\n2. [Données et illustrations] ............................ 5\nExercice 01 — [Titre] ...................................... 6\nAnnexes et références .................................... 7')
para('Actualisez le sommaire après avoir remplacé les titres et supprimé les pages inutiles.','3iL Repère')
heading('Objectifs du module')
for text in ['[Identifier / expliquer la notion…]','[Appliquer la méthode à une situation…]','[Produire et justifier un résultat…]']:para(text,'3iL Puces')
heading('Prérequis et matériel')
para('[Indiquez les connaissances nécessaires, les outils à installer et les ressources à prévoir.]')

title('1. [Titre du chapitre]')
callout('Objectif du chapitre','[À la fin de ce chapitre, vous serez capable de…]')
heading('1.1 [Titre de section]')
para('[Présentez la notion en quelques paragraphes. Conservez un vocabulaire précis et expliquez les termes nouveaux avant de les employer.]')
para('Définition — [Terme] : [définition courte et précise].','3iL Définition')
para('Exemple — [Présentez un cas concret qui illustre la notion.]','3iL Exemple')
para('Point de vigilance — [Indiquez une erreur fréquente et la façon de l’éviter.]','3iL Attention')
heading('1.2 [Méthode à appliquer]')
for text in ['[Préparez les éléments nécessaires.]','[Réalisez l’opération ou le raisonnement.]','[Vérifiez le résultat obtenu.]']:para(text,'3iL Étapes')
para('1.2.1 [Précision utile]','Heading 3')
para('[Ajoutez ici un complément si la notion le nécessite.]')
callout('À retenir','[Formulez l’idée essentielle du chapitre en une ou deux phrases.]')

title('2. [Données et illustrations]')
heading('2.1 [Titre du tableau]')
tab=d.add_table(rows=4,cols=3);tab.autofit=False
widths=[2400,3600,3638];pr=tab._tbl.tblPr;pr.find(qn('w:tblW')).set(qn('w:w'),str(sum(widths)));pr.find(qn('w:tblW')).set(qn('w:type'),'dxa');pr.append(el('tblInd',w=120,type='dxa'))
mar=el('tblCellMar')
for edge in ['top','bottom','left','right']:mar.append(el(edge,w=120,type='dxa'))
pr.append(mar)
for col,w in zip(tab._tbl.tblGrid,widths):col.set(qn('w:w'),str(w))
for i,row in enumerate(tab.rows):
    if i==0:row._tr.get_or_add_trPr().append(el('tblHeader'))
    for j,cell in enumerate(row.cells):
        cell.width=Pt(widths[j]/20);cell._tc.get_or_add_tcPr().append(el('vAlign',val='center'));cell._tc.get_or_add_tcPr().append(el('shd',fill=BLUE if i==0 else SOFT if i%2 else 'FFFFFF'))
        p=cell.paragraphs[0];p.style='3iL Tableau';r=p.add_run(['[Critère]','[Élément A]','[Élément B]'][j] if i==0 else '[À compléter]')
        if i==0:r.bold=True;r.font.color.rgb=RGBColor(255,255,255)
para('Tableau 1 — [Légende et source des données].','Caption')
heading('2.2 [Titre de la figure]')
p=para('[Insérez votre image ou votre schéma ici]','3iL Repère');p.paragraph_format.space_before=Pt(22);p.paragraph_format.space_after=Pt(22)
para('Figure 1 — [Ce que montre la figure]. Source : [auteur, date, référence].','Caption')
para('[Expliquez le lien entre la figure et la notion étudiée.]')
heading('2.3 [Extrait de code ou formule]')
para('[Première ligne de votre exemple]\n    [Ligne indentée]\n[Dernière ligne]','3iL Code')
para('[Expliquez le résultat attendu et les éléments importants de l’extrait.]')

title('Exercice 01 — [Titre]')
callout('Modalités','[Durée] · [Individuel / groupe] · [Livrable attendu]')
heading('Contexte et objectif')
para('[Décrivez la situation, les données disponibles et la compétence à mobiliser.]')
heading('Travail demandé')
for text in ['[Première consigne observable.]','[Deuxième consigne observable.]','[Justification ou vérification attendue.]']:para(text,'3iL Puces')
heading('Critères de réussite')
para('[Précisez ce qui permet d’évaluer la qualité de la réponse : exactitude, méthode, justification et présentation.]')
heading('Votre réponse')
para('[Rédigez votre réponse ou indiquez le support de restitution demandé.]')

title('Annexes et références')
heading('Annexe A — [Titre de la ressource]')
para('[Ajoutez les documents utiles : jeu de données, consigne détaillée, procédure ou approfondissement.]')
heading('Références du cours')
for text in ['[Auteur / organisme]. [Titre]. [Date ou version]. [URL ou référence bibliographique].','[Auteur / organisme]. [Titre]. [Date ou version]. [URL ou référence bibliographique].']:para(text,'3iL Puces')
heading('Crédits des illustrations')
para('[Figure concernée] — [Auteur, provenance, licence ou autorisation].')

for sec in d.sections:
    p=sec.header.paragraphs[0];p.clear();p.add_run('3iL PROGRAMMES EXPERTS  /  SUPPORT DE COURS')
    p=sec.footer.paragraphs[0];p.clear();p.add_run('© ');field(p,'DATE \\@ "yyyy"','2026');p.add_run(' 3iL Programmes Experts     /     ');field(p,'PAGE','1')
settings=d.settings.element
for x in list(settings.findall(qn('w:updateFields'))):settings.remove(x)
settings.append(el('updateFields',val='true'))
# Drop relationships to removed UML illustrations and links, not the official cover.
used=set(d._element.xpath('//@r:embed | //@r:id'))
for rid,rel in list(d.part.rels.items()):
    if rel.reltype.rsplit('/',1)[-1] in ['image','hyperlink'] and rid not in used:d.part.drop_rel(rid)
d.core_properties.title='Gabarit de cours — 3iL Programmes Experts'
d.core_properties.subject='Modèle pédagogique réutilisable';d.core_properties.author='3iL Programmes Experts'
d.core_properties.comments='';d.core_properties.keywords='gabarit, cours, 3iL'
dest=OUT/'3iL-Gabarit-Cours.docx';d.save(dest)
with zipfile.ZipFile(dest) as z,zipfile.ZipFile(OUT/'3iL-Gabarit-Cours.dotx','w',zipfile.ZIP_DEFLATED) as out:
    for name in z.namelist():
        data=z.read(name)
        if name=='[Content_Types].xml':data=data.replace(b'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml',b'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml')
        out.writestr(name,data)
assert sha256(REF.read_bytes()).hexdigest()==before
with zipfile.ZipFile(REF) as a,zipfile.ZipFile(dest) as b:
    source_doc=Document(REF)
    for name in ['Normal','Title','Subtitle','Heading 1','Heading 2','Heading 3','3iL Repère','3iL Code','Caption']:
        for tag in ['w:pPr','w:rPr']:
            x=source_doc.styles[name].element.find(qn(tag));y=d.styles[name].element.find(qn(tag))
            assert (x.xml if x is not None else None)==(y.xml if y is not None else None),(name,tag)
    inventory={n:sha256(a.read(n)).hexdigest() for n in a.namelist()}
(TMP/'word-source-evidence.json').write_text(json.dumps({'reference_sha256':before,'parts':inventory,'output':str(dest)},indent=2),encoding='utf8')
print(dest)
