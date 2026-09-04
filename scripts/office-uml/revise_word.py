"""Local edits to the existing course; originals and every exercise are preserved."""
from pathlib import Path
import json, re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from word_editorial import apply_editorial_changes

ROOT=Path(__file__).resolve().parents[2]
OUT=ROOT/'office/UML'
manifest=json.loads((OUT/'diagrammes/manifest.json').read_text(encoding='utf-8'))['figures']
doc=Document(OUT/'3iL-UML-Cours-complet.docx')
removed=[]; count=0; groups=[]
for p in list(doc.paragraphs):
    descriptors=p._p.xpath('.//wp:docPr/@descr')
    if not descriptors: continue
    key=descriptors[0]
    if key=='Diagramme de composants': key='07-example'
    if key=='Diagramme de déploiement': key='08-example'
    if key not in manifest: continue
    groups.append(key)
    following=p._p.getnext()
    if following is not None and ''.join(following.itertext()).startswith('Figure'):
        following.getparent().remove(following)
    for part in manifest[key]:
        count+=1
        img=OUT/'diagrammes'/f"{part['id']}.png"
        w,h=Image.open(img).size
        scale=min(17/w,18/h)
        new=p.insert_paragraph_before()
        new.alignment=WD_ALIGN_PARAGRAPH.CENTER
        new.paragraph_format.keep_with_next=True
        new.paragraph_format.space_before=Pt(10)
        shape=new.add_run().add_picture(str(img),width=Cm(w*scale),height=Cm(h*scale))
        shape._inline.docPr.set('descr',part['title'])
        cap=p.insert_paragraph_before(f"Figure {count} — {part['title']}",'Caption')
        cap.paragraph_format.keep_with_next=False
    p._p.getparent().remove(p._p)
for p in list(doc.paragraphs):
    if p.text.startswith(('Rectification :','Lecture critique :')):
        removed.append(p.text);p._p.getparent().remove(p._p)
    elif p.text.startswith('Les coquilles et plusieurs erreurs de fond ont été corrigées :'):
        removed.append(p.text)
        p.text='Les diagrammes emploient une notation homogène. Dans le cas de la commande de repas, préparation et validation du paiement progressent en parallèle ; la livraison attend leur achèvement. Les exemples de modèles de données complètent UML sans appartenir à sa classification.'
assert len(groups)==31,groups
assert count==sum(map(len,manifest.values()))
apply_editorial_changes(doc)
doc.save(OUT/'3iL-UML-Cours-complet-revise.docx')
(ROOT/'tmp/office-uml/word-revision-log.json').write_text(json.dumps({'groups':groups,'panels':count,'removed':removed},ensure_ascii=False,indent=2),encoding='utf-8')
print(f'Word revised: {len(groups)} figures / {count} panels')
