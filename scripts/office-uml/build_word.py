"""Build the editable 3iL course from the complete, normalized Markdown corpus."""
from pathlib import Path
import json, re, html, unicodedata, textwrap
from urllib.parse import unquote
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

BASE=Path(__file__).resolve().parent
ROOT=BASE.parent.parent
OUT=ROOT/'office'/'UML'
OUT.mkdir(parents=True,exist_ok=True)
sources=json.loads((BASE/'sources.json').read_text(encoding='utf8'))
BLUE='005067'; INK='122F37'; MUTED='5D747A'; ORANGE='E84D0D'; SOFT='EDF4F5'
doc=Document(); sec=doc.sections[0]
for s in doc.styles:
    for border in s.element.findall('.//'+qn('w:pBdr')):border.getparent().remove(border)
sec.page_width=Cm(21);sec.page_height=Cm(29.7)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Cm(2)
sec.header_distance=sec.footer_distance=Cm(.8)
sec.different_first_page_header_footer=True

def el(tag,**attrs):
    e=OxmlElement('w:'+tag)
    for k,v in attrs.items():e.set(qn('w:'+k),str(v))
    return e

def style(name,size,color=INK,bold=False,before=0,after=6,keep=False,font='Arial'):
    s=doc.styles[name] if name in doc.styles else doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
    s.font.name=font;s.font.size=Pt(size);s.font.bold=bold;s.font.italic=False;s.font.color.rgb=RGBColor.from_string(color)
    p=s.paragraph_format;p.space_before=Pt(before);p.space_after=Pt(after);p.line_spacing=1.15
    p.keep_with_next=keep;p.widow_control=True
    s.element.get_or_add_rPr().append(el('lang',val='fr-FR'))
    return s
style('Normal',10.5)
style('Title',40,BLUE,True,0,14,True)
style('Subtitle',18,MUTED,False,0,12)
for i,size in enumerate([24,16,12,11],1):style('Heading '+str(i),size,BLUE,True,16 if i>1 else 0,8,True)
style('3iL Chapitre',10,ORANGE,True,0,10,True)
style('3iL Repère',10,BLUE,False,8,8)
style('3iL Code',8.5,INK,False,0,0,font='Consolas').paragraph_format.line_spacing=1.05
style('Caption',9,MUTED,False,4,10)
style('3iL Navigation',11,BLUE,False,0,7)
style('3iL Liste',10.5,INK,False,0,4)
style('Header',8,MUTED,False,0,0)
style('Footer',8,MUTED,False,0,0)
style('3iL Tableau',9.5,INK,False,0,4)
for name in ['3iL Repère','3iL Code']:
    pp=doc.styles[name].element.get_or_add_pPr();pp.append(el('shd',fill=SOFT))
    if name=='3iL Repère':
        border=el('pBdr');border.append(el('left',val='single',sz=18,space=8,color=BLUE));pp.append(border)
        doc.styles[name].paragraph_format.left_indent=Cm(.35)

def field(p,code,text=''):
    r=p.add_run();r._r.append(el('fldChar',fldCharType='begin'))
    t=el('instrText');t.set(qn('xml:space'),'preserve');t.text=' '+code+' ';r._r.append(t)
    r._r.append(el('fldChar',fldCharType='separate'));p.add_run(text)
    p.add_run()._r.append(el('fldChar',fldCharType='end'))

hp=sec.header.paragraphs[0];hp.style='Header'
hp.add_run('3iL PROGRAMMES EXPERTS   /   UML')
hp.paragraph_format.tab_stops.add_tab_stop(Cm(17),WD_ALIGN_PARAGRAPH.RIGHT)
fp=sec.footer.paragraphs[0];fp.style='Footer'
fp.add_run('© 2026 3iL Programmes Experts');fp.add_run(' '*5+'UML  /  ');field(fp,'PAGE','1')
fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT

def plain(tokens):
    out=''
    for t in tokens:
        typ=t['type']
        if typ=='image':continue
        if typ=='br':out+='\n'
        elif 'tokens'in t:out+=plain(t['tokens'])
        else:out+=html.unescape(t.get('text',''))
    return out

anchors={unicodedata.normalize('NFC',s['path']):'source_'+str(i) for i,s in enumerate(sources)}
bookmark_id=0
def bookmark(p,name):
    global bookmark_id
    bookmark_id+=1
    p._p.insert(0,el('bookmarkStart',id=bookmark_id,name=name));p._p.append(el('bookmarkEnd',id=bookmark_id))

def link(p,text,href,current=''):
    h=el('hyperlink')
    if href.startswith('http'):
        h.set(qn('r:id'),p.part.relate_to(href,RT.HYPERLINK,is_external=True))
    else:
        target=(BASE/'source'/Path(current).parent/unquote(href)).resolve()
        key=unicodedata.normalize('NFC',target.relative_to(BASE/'source').as_posix())
        if key not in anchors:raise ValueError('Unresolved source link: '+key)
        h.set(qn('w:anchor'),anchors[key])
    r=el('r');pr=el('rPr');pr.append(el('color',val=BLUE));pr.append(el('u',val='single'));r.append(pr)
    t=el('t');t.text=text;r.append(t);h.append(r);p._p.append(h)

figure_notes={
 '01-hierarchy.png':'Lecture critique : cette illustration source omet les profils et classe mal les diagrammes d’interaction. Séquence, communication, timing et vue d’ensemble d’interaction sont des diagrammes de comportement ; objets, classes et profils sont structurels.',
 '02-symbols.png':'Rectification : la généralisation se termine par un triangle creux orienté vers l’acteur général, et non une pointe pleine.',
 '02-example.png':'Lecture critique : les relations include/extend du schéma source ne constituent pas un corrigé. Reliez des cas d’utilisation, jamais des cadres. « include » vise le cas inclus ; « extend » vise le cas étendu. Justifiez chaque inclusion ou extension par le scénario.',
 '04-example.png':'Lecture critique : une croix indique une destruction, pas simplement la fin du scénario. Le client et les acteurs persistants ne sont pas détruits à la fin d’une commande.',
 '04-symbols.png':'Rectification : « retour asynchrone » n’est pas un type UML distinct. Une réponse applicative asynchrone est un nouveau message.',
 '05-symbols.png':'Rectification : écrivez les multiplicités UML 1..* et 0..*, et non 1,* ou 0,n. Dérivé et statique ne sont pas des visibilités.',
 '06a-example.png':'Rectification : identifiez et soulignez les instances, par exemple client1:Client. Les liens d’instances représentent des associations concrètes ; les flèches de messages appartiennent à une vue d’interaction.',
 '06b-example.png':'Rectification : l’ordre d’identification est nomInstance:Classe, par exemple sophie:Client, et non Client:Sophie Dupont.',
 '06c-example.png':'Lecture critique : placez une garde sur une transition sous la forme événement [condition] / effet ; ne confondez pas état et action. Le parallélisme d’états demande des régions orthogonales.',
 '06c-symbols.png':'Rectification : le losange représente un choix ; la garde est l’expression entre crochets portée par la transition.',
 '09e-example.png':'Repère : entité-association et entité-relation désignent la même famille de modélisation conceptuelle ; les notations Merise, Chen et patte d’oie diffèrent.',
 '09f-symbols.png':'Repère : la patte d’oie est une notation entité-relation, pas la notation originale de Chen.',
 '09h-example.png':'Repère : familles de colonnes NoSQL et stockage colonnaire analytique sont deux notions distinctes. Les garanties et performances dépendent du produit et des requêtes.'
}
used_images=set();figures=0
def add_image(href,current,alt=None):
    global figures
    path=BASE/'media'/Path(unquote(href)).name
    used_images.add(path.name);figures+=1
    w,h=Image.open(path).size
    maxw=17;maxh=18.2 if path.name in figure_notes else 20.2
    scale=min(maxw/w,maxh/h)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.keep_with_next=True
    pic=p.add_run().add_picture(str(path),width=Cm(w*scale),height=Cm(h*scale))
    pic._inline.docPr.set('descr',alt or path.stem)
    cap=doc.add_paragraph(f'Figure {figures} — '+(alt or path.stem)+' · Illustration du dossier source','Caption')
    if path.name in figure_notes:
        cap.paragraph_format.keep_with_next=True
        doc.add_paragraph(figure_notes[path.name],'3iL Repère')

def inline(p,tokens,current,bold=False,italic=False):
    for t in tokens:
        typ=t['type']
        if typ=='image':add_image(t['href'],current)
        elif typ=='link':link(p,plain(t.get('tokens',[])) or t.get('text',''),t['href'],current)
        elif typ in ('strong','em'):inline(p,t['tokens'],current,bold or typ=='strong',italic or typ=='em')
        elif typ in ('text','escape','codespan'):
            if 'tokens' in t:inline(p,t['tokens'],current,bold,italic);continue
            r=p.add_run(html.unescape(t.get('text','')).replace('\n',' '));r.bold=True if bold else None;r.italic=True if italic else None
            if typ=='codespan':r.font.name='Consolas';r.font.size=Pt(9)
        elif typ=='br':p.add_run().add_break()
        elif typ=='html' and re.match(r'<br\s*/?>',t.get('text','')):p.add_run().add_break()
        elif typ=='del':inline(p,t['tokens'],current,bold,italic)

def number_list(ordered,start=1):
    numroot=doc.part.numbering_part.element
    aid=max([int(x.get(qn('w:abstractNumId'))) for x in numroot.findall(qn('w:abstractNum'))]+[0])+1
    a=el('abstractNum',abstractNumId=aid);a.append(el('multiLevelType',val='multilevel'))
    for level in range(9):
        lv=el('lvl',ilvl=level);lv.append(el('start',val=start if level==0 else 1));lv.append(el('numFmt',val='decimal' if ordered else 'bullet'))
        lv.append(el('lvlText',val=f'%{level+1}.' if ordered else '•'));lv.append(el('lvlJc',val='left'))
        pp=el('pPr');pp.append(el('ind',left=360*(level+1),hanging=240));lv.append(pp);a.append(lv)
    numroot.append(a)
    nid=max([int(x.get(qn('w:numId'))) for x in numroot.findall(qn('w:num'))]+[0])+1
    n=el('num',numId=nid);n.append(el('abstractNumId',val=aid));numroot.append(n);return nid

def list_block(t,current,level=0):
    nid=number_list(t.get('ordered',False),t.get('start') or 1)
    for item in t['items']:
        first=True
        for b in item.get('tokens',[]):
            if b['type']=='list':list_block(b,current,level+1)
            elif b['type'] in ('text','paragraph'):
                p=doc.add_paragraph(style='3iL Liste');pp=p._p.get_or_add_pPr()
                if first:
                    np=el('numPr');np.append(el('ilvl',val=level));np.append(el('numId',val=nid));pp.append(np)
                else:p.paragraph_format.left_indent=Pt(18*(level+1))
                inline(p,b.get('tokens',[{'type':'text','text':b.get('text','')}]),current);first=False
            else:blocks([b],current)

def table_block(t,current):
    rows=[t['header']]+t['rows'];cols=len(rows[0]);widths=[1700,3800,4138] if cols==3 else [9638//cols]*cols
    tab=doc.add_table(rows=len(rows),cols=cols);tab.autofit=False
    pr=tab._tbl.tblPr;pr.find(qn('w:tblW')).set(qn('w:w'),str(sum(widths)));pr.find(qn('w:tblW')).set(qn('w:type'),'dxa');pr.append(el('tblInd',w=120,type='dxa'))
    mar=el('tblCellMar')
    for edge in ['top','bottom','left','right']:mar.append(el(edge,w=120,type='dxa'))
    pr.append(mar)
    for col,w in zip(tab._tbl.tblGrid,widths):col.set(qn('w:w'),str(w))
    for i,row in enumerate(rows):
        if i==0:tab.rows[i]._tr.get_or_add_trPr().append(el('tblHeader'))
        for j,c in enumerate(row):
            cell=tab.cell(i,j);cell.width=Pt(widths[j]/20);cell._tc.get_or_add_tcPr().append(el('vAlign',val='center'))
            cell._tc.get_or_add_tcPr().append(el('shd',fill=BLUE if i==0 else SOFT if i%2 else 'FFFFFF'))
            p=cell.paragraphs[0];p.style='3iL Tableau';inline(p,c.get('tokens',[]),current)
            if i==0:
                for r in p.runs:r.bold=True;r.font.color.rgb=RGBColor(255,255,255)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def blocks(tokens,current,skip_title=False):
    for i,t in enumerate(tokens):
        typ=t['type']
        if typ=='space':continue
        if typ=='heading':
            if skip_title and i==0:continue
            txt=plain(t.get('tokens',[]));d=min(t['depth'],4)
            p=doc.add_paragraph(style='Heading '+str(d));inline(p,t['tokens'],current)
        elif typ in ('paragraph','text'):
            children=t.get('tokens',[])
            if children and all(x['type'] in ('image','br') for x in children):
                for x in children:
                    if x['type']=='image':add_image(x['href'],current)
            else:inline(doc.add_paragraph(),children,current)
        elif typ=='list':list_block(t,current)
        elif typ=='code':
            doc.add_paragraph((t.get('lang') or 'Code').upper(),'3iL Chapitre')
            lines=t['text'].splitlines()
            for line in lines:
                # Visual wraps preserve the exact source characters and indentation.
                p=doc.add_paragraph(line,'3iL Code');p.paragraph_format.keep_with_next=False
            doc.add_paragraph().paragraph_format.space_after=Pt(3)
        elif typ=='blockquote':
            p=doc.add_paragraph(plain(t.get('tokens',[])).strip(),'3iL Repère')
        elif typ=='table':table_block(t,current)
        elif typ=='hr':doc.add_paragraph()
        elif typ=='html':
            txt=re.sub('<[^>]+>','',t.get('text','')).strip()
            if txt:doc.add_paragraph(html.unescape(txt))
        else:raise ValueError('Unhandled block: '+typ)

# Editorial cover: simple, native Word paragraphs, no layout tables.
p=doc.add_paragraph();p.add_run().add_picture(str(ROOT/'assets/logos/3il-horizontal.png'),width=Cm(10))
p.paragraph_format.space_after=Pt(80)
doc.add_paragraph('SUPPORT DE COURS','3iL Chapitre')
doc.add_paragraph('UML', 'Title')
doc.add_paragraph('Modéliser les systèmes\net les données','Subtitle')
doc.add_paragraph('Cours complet · 12 exercices · 2 annexes','3iL Repère')
doc.add_paragraph('Programmation orientée objet\nBases de données SQL et NoSQL')
doc.add_paragraph('3iL Programmes Experts','Heading 2')
doc.add_page_break()
p=doc.add_paragraph('Parcours du cours','Heading 1');bookmark(p,'sommaire')
doc.add_paragraph('Les chapitres conservent les repères du dossier d’origine. Les exercices sont regroupés après le cours ; les deux annexes suivent les exercices. Les liens ci-dessous permettent d’y accéder directement.')
courses=[s for s in sources if s['path'].startswith('Cours/') and not s['name'].startswith('Annexe')]
exercises=[s for s in sources if s['path'].startswith('Exercices/')]
annexes=[s for s in sources if s['name'].startswith('Annexe')]
order=courses+exercises+annexes
for label,group in [('Cours',courses),('Exercices',exercises),('Annexes',annexes)]:
    doc.add_paragraph(label,'Heading 2')
    for s in group:
        p=doc.add_paragraph(style='3iL Navigation');link(p,s['name'][:-3],s['path'],'index.md')
for s in order:
    p=doc.add_paragraph('EXERCICE' if s in exercises else 'ANNEXE' if s in annexes else 'COURS','3iL Chapitre')
    p.paragraph_format.page_break_before=True
    bookmark(p,anchors[unicodedata.normalize('NFC',s['path'])])
    title=plain(s['tokens'][0].get('tokens',[]))
    if s in exercises:title=title.replace(' - Exercice : ','\n')
    if s in annexes:title=s['name'][:-3].replace(' - ','\n',1)
    doc.add_paragraph(title,'Heading 1')
    blocks(s['tokens'],s['path'],True)
p=doc.add_paragraph('Illustrations complémentaires','Heading 1');p.paragraph_format.page_break_before=True
doc.add_paragraph('Ces deux illustrations sont présentes dans le dossier source, sans chapitre Markdown dédié. Elles complètent les vues structurelles présentées en introduction.')
for f,title in [('07-example.png','Diagramme de composants'),('08-example.png','Diagramme de déploiement')]:
    doc.add_paragraph(title,'Heading 2');add_image('../images/'+f,courses[0]['path'],title)
p=doc.add_paragraph('Repères de lecture et sources','Heading 1');p.paragraph_format.page_break_before=True
doc.add_paragraph('Ce support reprend l’ensemble des 21 fichiers Markdown du dossier « 01 - UML », avec leurs exercices, annexes, exemples de code et illustrations. Les PDF présents dans le dossier reprennent ces contenus et ne sont pas reproduits une seconde fois.')
doc.add_paragraph('Les coquilles et plusieurs erreurs de fond ont été corrigées : substitution de Liskov, exemple Animal, ligne de vie, vocabulaire des diagrammes, numérotation et distinction des modèles de données. Les illustrations originales présentant des défauts sont explicitement accompagnées d’une rectification ou d’une consigne de lecture critique. Elles ne doivent pas être utilisées isolément comme corrigés.','3iL Repère')
for title,url in [
 ('OMG — UML 2.5.1','https://www.omg.org/spec/UML/2.5.1'),
 ('Liskov et Wing — A Behavioral Notion of Subtyping','https://www.cs.cmu.edu/afs/cs/project/venari/www/subtype-toplas.html'),
 ('Apache Cassandra — conception du modèle de données','https://cassandra.apache.org/doc/stable/cassandra/developing/data-modeling/data-modeling_rdbms.html'),
 ('MongoDB — validation du schéma','https://www.mongodb.com/docs/manual/core/schema-validation/'),
 ('MongoDB — transactions','https://www.mongodb.com/docs/manual/data-modeling/enforce-consistency/transactions/')]:
    p=doc.add_paragraph();link(p,title,url)
doc.core_properties.title='UML — Cours, exercices et annexes'
doc.core_properties.subject='3iL Programmes Experts — Support pédagogique'
doc.core_properties.author='3iL Programmes Experts';doc.core_properties.last_modified_by='3iL Programmes Experts'
doc.core_properties.comments=''
out=OUT/'3iL-UML-Cours-complet.docx';doc.save(out)
stats={'files':len(order),'course_files':len(courses),'exercises':len(exercises),'annexes':len(annexes),'figures':figures,'images':sorted(used_images),'paragraphs':len(doc.paragraphs),'tables':len(doc.tables)}
(BASE/'word-audit.json').write_text(json.dumps(stats,indent=2,ensure_ascii=False),encoding='utf8')
print(json.dumps(stats,ensure_ascii=True))
