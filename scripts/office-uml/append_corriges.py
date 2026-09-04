"""Append a separate answer appendix; leave all original exercise bodies intact."""
from pathlib import Path
import json, hashlib
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT=Path(__file__).resolve().parents[2]
WORK=ROOT/'tmp/uml-corriges';WORK.mkdir(exist_ok=True,parents=True)
src=ROOT/'office/UML/3iL-UML-Cours-complet-revise.docx'
if not src.exists(): src=ROOT/'office/UML/3iL-UML-Cours.docx'
d=Document(src)
existing=next((p for p in d.paragraphs if p.text=='Annexe 03 — Corrigés des exercices' and p.style.name=='Heading 1'),None)
if existing is not None:
    # Regenerate only this generated appendix from a final document copy.
    node=existing._p.getprevious()
    assert node is not None and node.xpath('.//w:br[@w:type="page"]')
    while node is not None and node.tag!=qn('w:sectPr'):
        following=node.getnext();node.getparent().remove(node);node=following
    for p in list(d.paragraphs):
        if p._p.xpath('.//w:hyperlink[@w:anchor="corriges"]'):
            p._p.getparent().remove(p._p)
original=[str(p._p.xml) for p in d.paragraphs]
specs=json.loads((ROOT/'scripts/office-uml/corriges_exercices.json').read_text(encoding='utf-8'))
def mark(p,name,number):
    start=OxmlElement('w:bookmarkStart');start.set(qn('w:id'),str(number));start.set(qn('w:name'),name)
    end=OxmlElement('w:bookmarkEnd');end.set(qn('w:id'),str(number));p._p.insert(0,start);p._p.append(end)
def link(p,text,anchor):
    h=OxmlElement('w:hyperlink');h.set(qn('w:anchor'),anchor)
    r=OxmlElement('w:r');pr=OxmlElement('w:rPr');c=OxmlElement('w:color');c.set(qn('w:val'),'005067');pr.append(c)
    u=OxmlElement('w:u');u.set(qn('w:val'),'single');pr.append(u);r.append(pr)
    t=OxmlElement('w:t');t.text=text;r.append(t);h.append(r);p._p.append(h)

# Minimal front matter updates; original chapters and exercises remain identical.
d.paragraphs[4].text='Cours complet · 12 exercices · 3 annexes dont les corrigés'
d.paragraphs[9].text='Les exercices sont regroupés après le cours. Les annexes de POO et de bases de données suivent ; les corrigés des 12 exercices sont réunis dans une annexe séparée en fin de document.'
p=d.paragraphs[34].insert_paragraph_before(style='3iL Navigation');link(p,'Annexe 03 — Corrigés des exercices','corriges')
d.add_page_break()
p=d.add_paragraph('Annexe 03 — Corrigés des exercices','Heading 1');mark(p,'corriges',5000)
d.add_paragraph('Propositions de résolution','Heading 2')
d.add_paragraph('Cette annexe regroupe les réponses aux 12 exercices, dans le même ordre que les énoncés. Elle est séparée des exercices pour permettre de chercher avant de consulter la solution. Les modèles proposés ne sont pas les seules réponses possibles : une variante cohérente et justifiée peut être recevable.')
d.add_paragraph('Les hypothèses complètent uniquement les points laissés ouverts par les énoncés. Les figures de structure montrent les associations principales ; les attributs, opérations et contraintes sont détaillés dans le texte. Les deux figures de séquence illustrent le scénario nominal : les alternatives font partie du corrigé écrit.','3iL Repère')
for i,spec in enumerate(specs,1):
    p=d.add_paragraph(style='3iL Navigation');link(p,f'{i:02} — '+spec['title'],f'corrige_{i}')
d.add_paragraph('Référence de notation : OMG, Unified Modeling Language 2.5.1, https://www.omg.org/spec/UML/2.5.1. Les modèles entité-relation sont des compléments de modélisation des données, distincts de la classification UML.','Caption')
for i,spec in enumerate(specs,1):
    d.add_page_break();p=d.add_paragraph(f'Corrigé {i:02}\n'+spec['title'],'Heading 1');mark(p,f'corrige_{i}',5000+i)
    p=d.add_paragraph(style='3iL Navigation');link(p,'Revenir à l’énoncé',f"source_{spec['source']}")
    for heading,body in spec['sections']:
        d.add_paragraph(heading,'Heading 3');d.add_paragraph(body)
    d.add_page_break();d.add_paragraph(f'Corrigé {i:02} — Proposition de schéma','Heading 2')
    path=WORK/'figures'/f"{spec['diagram']}.png";w,h=Image.open(path).size
    # A full portrait page for each figure keeps complex models legible.
    scale=min(17/w,21/h)
    p=d.add_paragraph();p.paragraph_format.keep_with_next=True
    pic=p.add_run().add_picture(str(path),width=Cm(w*scale),height=Cm(h*scale));pic._inline.docPr.set('descr',spec['title'])
    d.add_paragraph(f'Schéma C{i:02} — '+spec['title']+' · À lire avec les hypothèses et contraintes du corrigé.','Caption')
out=WORK/'3iL-UML-Cours.docx';d.save(out)
after=Document(out)
# All old paragraphs except the two explicit front-matter edits survive byte-for-byte.
updated={str(p._p.xml) for p in after.paragraphs}
assert all(xml in updated for i,xml in enumerate(original) if i not in (4,9))
anchors={e.get(qn('w:name')) for e in after.element.xpath('.//w:bookmarkStart')}
assert all(e.get(qn('w:anchor')) in anchors for e in after.element.xpath('.//w:hyperlink[@w:anchor]'))
assert len(specs)==12
(WORK/'audit.json').write_text(json.dumps({'source_sha256':hashlib.sha256(src.read_bytes()).hexdigest(),'corriges':12,'diagrams':12,'original_paragraphs_preserved':len(original)-2,'internal_links':'valid'},indent=2),encoding='utf-8')
print(out)
