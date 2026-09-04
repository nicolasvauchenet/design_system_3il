"""Read-only checks of the revised delivery, coverage and preserved course content."""
from pathlib import Path
import json,re,zipfile
from lxml import etree
from docx import Document
from word_editorial import UPDATE, HISTORY_PREFIX, HISTORY_END, HISTORY_REPLACEMENT, SUFFIX, is_exercise_reference

ROOT=Path(__file__).resolve().parents[2]
OUT=ROOT/'office/UML'
ns={'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main','a':'http://schemas.openxmlformats.org/drawingml/2006/main'}
docs=json.loads((ROOT/'scripts/office-uml/sources.json').read_text(encoding='utf8'))
manifest=json.loads((OUT/'diagrammes/manifest.json').read_text(encoding='utf8'))['figures']
panels=[p for group in manifest.values() for p in group]
assert len(manifest)==31 and len(panels)==52
for p in panels:
    for ext in ['dot','svg','png']:assert (OUT/'diagrammes'/f"{p['id']}.{ext}").stat().st_size>0
with zipfile.ZipFile(OUT/'3iL-UML-Cours-complet-revise.docx') as z:
    for f in z.namelist():
        if f.endswith('.xml'):etree.fromstring(z.read(f))
    x=etree.fromstring(z.read('word/document.xml'))
    names=set(x.xpath('//w:bookmarkStart/@w:name',namespaces=ns))
    anchors=x.xpath('//w:hyperlink/@w:anchor',namespaces=ns)
    assert all(a in names for a in anchors)
    word='\n'.join(x.xpath('//w:t/text()',namespaces=ns))
    for forbidden in ['Rectification :','Lecture critique :','Les coquilles et plusieurs erreurs de fond']:
        assert forbidden not in word,forbidden
    caps=re.findall(r'Figure (\d+) —',word)
    assert caps==list(map(str,range(1,53))),caps
old=Document(OUT/'3iL-UML-Cours-complet.docx')
new=Document(OUT/'3iL-UML-Cours-complet-revise.docx')
preserved=[]
for p in old.paragraphs:
    t=p.text
    if not t or t == UPDATE or t.startswith(('Figure ','Rectification :','Lecture critique :','Les coquilles et plusieurs erreurs de fond')):
        continue
    if t.startswith(HISTORY_PREFIX):
        t=t[:-len(HISTORY_END)]+HISTORY_REPLACEMENT
    if is_exercise_reference(p):
        t+=SUFFIX
    preserved.append(t)
newtext=[p.text for p in new.paragraphs]
assert all(t in newtext for t in preserved),'A non-figure course paragraph is missing'
assert UPDATE not in newtext
assert sum(is_exercise_reference(p) and p.text.endswith(SUFFIX) for p in new.paragraphs)==12
brand=next(p for p in new.paragraphs if p.text=='3iL Programmes Experts')
assert brand.style.name=='3iL Établissement'
assert brand.style.element.xpath('./w:pPr/w:outlineLvl/@w:val')==['9']
assert [[c.text for row in t.rows for c in row.cells] for t in old.tables]==[[c.text for row in t.rows for c in row.cells] for t in new.tables]
with zipfile.ZipFile(OUT/'3iL-UML-Presentation-revisee.pptx') as z:
    slides=[];notes=[]
    for f in z.namelist():
        if f.endswith('.xml'):
            x=etree.fromstring(z.read(f))
            txt='\n'.join(x.xpath('//a:t/text()',namespaces=ns))
            if re.fullmatch(r'ppt/slides/slide\d+.xml',f):
                slides.append(f)
                assert not any(t in txt for t in ['Rectification','Lecture critique','Schéma source à discuter'])
            if re.fullmatch(r'ppt/notesSlides/notesSlide\d+.xml',f):notes.append(txt)
    norm=lambda s:re.sub(r'\s+',' ',s).strip()
    missing=[d['path'] for d in docs if not any(norm(d['raw']) in norm(n) for n in notes)]
    assert not missing,missing
    assert len(slides)==99,len(slides)
result={'figure_groups':31,'panels':52,'slides':len(slides),'complete_documents_in_notes':len(docs),'word_internal_links':len(anchors),'preserved_non_figure_paragraphs':len(preserved),'tables':'preserved','xml':'valid'}
print(json.dumps(result,ensure_ascii=False))
(ROOT/'tmp/office-uml/revision-check.json').write_text(json.dumps(result,indent=2),encoding='utf8')
